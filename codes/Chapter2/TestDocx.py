# 2장 처리 전 텍스트, 소싱, 정규화 - 파이썬에서 워드 문서 읽기
import docx
import word


# 문서 경로 초기화 및 전체 문서 출력
docName = '/workspace/NLP_python/codes/Chapter2/Files/sample-one-line.docx'
print('Document in full :\n', word.getTextWord(docName))

# 단락 정보
# 단락 수 출력
doc = docx.Document(docName)
print('단락 수 : ', len(doc.paragraphs))
# 2번 단락만 출력 및 2번 단락 스타일 출력
print('2번 단락 : ', doc.paragraphs[1].text)
print('2번 단락 스타일 : ', doc.paragraphs[1].style)

# 실행(run) 횟수 확인
print('Paragraph 1:', doc.paragraphs[0].text)
print('Number of runs in paragraph 1:', len(doc.paragraphs[0].runs))
for idx, run in enumerate(doc.paragraphs[0].runs):
    print('Run %s : %s' % (idx, run.text))

# 각 실행의 스타일 식별
print('is Run 5 underlined:', doc.paragraphs[0].runs[5].underline)
print('is Run 1 bold:', doc.paragraphs[0].runs[1].bold)
print('is Run 3 italic:', doc.paragraphs[0].runs[3].italic)
